import ttkbootstrap as ttk
from ttkbootstrap import Style
from tkinter import messagebox
from tkinter.filedialog import asksaveasfilename, askopenfilename
import tkinter as tk
import tkinter.font as tkFont  # Import the font module
from docx import Document
from docx.shared import Pt, Inches, RGBColor
import sqlite3
import json

# Application Defaults
THEME = (
    "darkly"  # cosmo, litera, minty, pulse, sandstone, solar, superhero, flatly, darkly
)
VERSION = "0.9"
DB_NAME = "outline.db"  # default db it will look for or create
GLOBAL_FONT_FAMILY = "Helvetica"  # Set the global font family
GLOBAL_FONT_SIZE = 12  # Set the global font size
GLOBAL_FONT = (GLOBAL_FONT_FAMILY, GLOBAL_FONT_SIZE)

# DOCX Exports
DOC_FONT = "Helvetica"
H1_SIZE = 18
H2_SIZE = 15
H3_SIZE = 12
H4_SIZE = 10
P_SIZE = 10
INDENT_SIZE = 0.25


class OutLineEditorApp:
    def __init__(self, root):
        # Apply ttkbootstrap theme
        self.style = Style(THEME)  # Use the global THEME
        self.root = root
        self.root.title(f"Outline Editor v{VERSION}")  # Use the global VERSION

        # Set global font scaling using tkinter.font
        default_font = tkFont.nametofont("TkDefaultFont")
        default_font.configure(family=GLOBAL_FONT_FAMILY, size=GLOBAL_FONT_SIZE)

        # Apply the font to Treeview specifically
        tree_font = tkFont.nametofont("TkTextFont")
        tree_font.configure(family=GLOBAL_FONT_FAMILY, size=GLOBAL_FONT_SIZE)

        # Apply the font to Label widgets
        label_font = tkFont.nametofont("TkHeadingFont")
        label_font.configure(family=GLOBAL_FONT_FAMILY, size=GLOBAL_FONT_SIZE)

        # Enforce Treeview row height based on font size
        row_height = int(GLOBAL_FONT_SIZE * 2.2)
        self.root.tk.call(
            "ttk::style", "configure", "Treeview", "-rowheight", row_height
        )

        # Padding constants
        LABEL_PADX = 5
        LABEL_PADY = (5, 5)  # Top padding: 5, bottom padding: 0
        ENTRY_PADY = (5, 5)  # Top padding: 0, bottom padding: 5
        SECTION_PADY = (5, 10)  # Treeview and Notes padding
        BUTTON_PADX = 5
        BUTTON_PADY = (5, 0)  # Button sections padding
        FRAME_PADX = 10
        FRAME_PADY = 10

        # Configure Treeview Frame
        self.tree_frame = ttk.Frame(root, width=500)  # Frame for Treeview
        self.tree_frame.grid(
            row=0, column=0, sticky="nswe", padx=FRAME_PADX, pady=FRAME_PADY
        )
        self.tree_frame.grid_propagate(False)  # Prevent shrinking
        self.tree_frame.grid_rowconfigure(1, weight=1)  # Let treeview expand vertically
        self.tree_frame.grid_columnconfigure(
            0, weight=1
        )  # Allow horizontal expansion within the frame

        # Add Label above Treeview
        ttk.Label(self.tree_frame, text="Your Outline", bootstyle="info").grid(
            row=0, column=0, sticky="w", padx=LABEL_PADX, pady=LABEL_PADY
        )

        # Treeview
        self.tree = ttk.Treeview(
            self.tree_frame, show="tree", bootstyle="info", height=25
        )
        self.tree.grid(
            row=1, column=0, sticky="nsew", padx=LABEL_PADX, pady=SECTION_PADY
        )
        self.tree.bind("<<TreeviewSelect>>", self.load_selected)

        # Make sure the Treeview expands fully within its frame
        self.tree_frame.grid_rowconfigure(1, weight=1)
        self.tree_frame.grid_columnconfigure(0, weight=1)

        # Add Search Entry (Search Bar)
        ttk.Label(
            self.tree_frame, text="Type Exact Search <Enter>", bootstyle="info"
        ).grid(row=2, column=0, sticky="w", padx=LABEL_PADX, pady=LABEL_PADY)
        self.search_entry = ttk.Entry(self.tree_frame, bootstyle="info")
        self.search_entry.grid(
            row=3, column=0, sticky="ew", padx=LABEL_PADX, pady=ENTRY_PADY
        )
        self.search_entry.bind(
            "<Return>", self.execute_search
        )  # Bind Enter key to search

        # Configure Editor Frame
        self.editor_frame = ttk.Frame(root)
        self.editor_frame.grid(
            row=0, column=1, sticky="nswe", padx=FRAME_PADX, pady=FRAME_PADY
        )
        self.editor_frame.grid_rowconfigure(3, weight=1)  # Allow textarea to expand
        self.editor_frame.grid_columnconfigure(
            0, weight=1
        )  # Allow horizontal expansion

        # Add Editor Fields
        ttk.Label(self.editor_frame, text="Title", bootstyle="info").grid(
            row=0, column=0, sticky="w", padx=LABEL_PADX, pady=LABEL_PADY
        )
        self.title_entry = ttk.Entry(self.editor_frame, bootstyle="info")
        self.title_entry.grid(
            row=1, column=0, sticky="ew", padx=LABEL_PADX, pady=ENTRY_PADY
        )

        ttk.Label(
            self.editor_frame, text="Questions Notes and Details", bootstyle="info"
        ).grid(row=2, column=0, sticky="w", padx=LABEL_PADX, pady=LABEL_PADY)
        self.questions_text = tk.Text(self.editor_frame, height=15)
        self.questions_text.grid(
            row=3, column=0, sticky="nswe", padx=LABEL_PADX, pady=SECTION_PADY
        )

        # Create Unified Button Row
        self.outliner_buttons = ttk.Frame(root)
        self.outliner_buttons.grid(
            row=1, column=0, columnspan=2, sticky="ew", padx=FRAME_PADX, pady=FRAME_PADY
        )

        # Add Buttons to Unified Button Row
        for text, command, style in [
            ("H(1)", self.add_header, "primary"),
            ("H(2)", self.add_category, "primary"),
            ("H(3)", self.add_subcategory, "primary"),
            ("H(4)", self.add_subheader, "primary"),
            ("(j) ↑", self.move_up, "secondary"),
            ("(k) ↓", self.move_down, "secondary"),
            ("(D)elete", self.delete_selected, "danger"),
            ("Make DOCX", self.export_to_docx, "success"),
            ("Load JSON", self.load_from_json, "info"),
            ("Load DB", self.load_database_from_file, "info"),
            ("New DB", self.reset_database, "warning"),
        ]:
            ttk.Button(
                self.outliner_buttons, text=text, command=command, bootstyle=style
            ).pack(side=tk.LEFT, padx=BUTTON_PADX, pady=BUTTON_PADY)

        # Configure root grid
        self.root.grid_rowconfigure(0, weight=1)  # Allow row 0 to expand
        self.root.grid_rowconfigure(1, weight=0)  # Make button row static in height
        self.root.grid_columnconfigure(0, minsize=400)  # Set minimum width for Treeview
        self.root.grid_columnconfigure(
            1, weight=1
        )  # Editor frame takes remaining space

        # Set initial window size
        self.root.geometry("900x800")  # Width x Height
        self.root.wm_minsize(825, 600)  # Minimum width and height

        # Database setup
        self.conn = sqlite3.connect(DB_NAME)
        self.cursor = self.conn.cursor()  # Initialize cursor before database methods
        self.setup_database()
        self.initialize_placement()

        self.last_selected_item_id = None

        # Bind focus-out events for auto-saving
        self.title_entry.bind("<FocusOut>", lambda event: self.save_data())
        self.questions_text.bind("<FocusOut>", lambda event: self.save_data())

        # Preload Data
        self.load_from_database()

        # Key Bindings
        self.root.bind_all("<Control-D>", lambda event: self.delete_selected())
        self.root.bind_all("<Control-d>", lambda event: self.delete_selected())
        self.root.bind_all("<Control-j>", lambda event: self.move_up())
        self.root.bind_all("<Control-k>", lambda event: self.move_down())
        self.root.bind_all("<F2>", self.focus_title_entry)
        self.root.bind_all("<Control-Key-1>", lambda event: self.add_header())
        self.root.bind_all("<Control-Key-2>", lambda event: self.add_category())
        self.root.bind_all("<Control-Key-3>", lambda event: self.add_subcategory())
        self.root.bind_all("<Control-Key-4>", lambda event: self.add_subheader())

        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)

    def focus_title_entry(self, event):
        """Move focus to the title entry and position the cursor at the end."""
        self.title_entry.focus_set()  # Focus on the title entry
        self.title_entry.icursor(tk.END)  # Move the cursor to the end of the text

    def execute_search(self, event=None):
        """Filter treeview to show only items that match the search query."""
        query = self.search_entry.get().strip()
        if not query:
            self.load_from_database()  # Reset tree if query is empty
            return

        # Query database for matching titles or questions
        self.cursor.execute(
            """
            WITH RECURSIVE parents AS (
                SELECT id, parent_id, title, questions
                FROM sections
                WHERE title LIKE ? OR questions LIKE ?
                UNION
                SELECT s.id, s.parent_id, s.title, s.questions
                FROM sections s
                INNER JOIN parents p ON s.id = p.parent_id
            )
            SELECT id, parent_id
            FROM parents
            ORDER BY parent_id, id
        """,
            (f"%{query}%", f"%{query}%"),
        )
        matches = self.cursor.fetchall()

        ids_to_show = {row[0] for row in matches}
        parents_to_show = {row[1] for row in matches if row[1] is not None}

        # Generate numbering for all items
        numbering_dict = self.generate_numbering()

        # Clear and repopulate the treeview
        self.tree.delete(*self.tree.get_children())
        self.populate_filtered_tree(None, "", ids_to_show, parents_to_show)

        # Apply consistent numbering
        self.calculate_numbering(numbering_dict)

    def populate_filtered_tree(
        self, parent_id, parent_node, ids_to_show, parents_to_show
    ):
        """Recursively populate the treeview with filtered data."""
        # Query children for the current parent_id
        if parent_id is None:
            self.cursor.execute(
                "SELECT id, title, parent_id FROM sections WHERE parent_id IS NULL"
            )
        else:
            self.cursor.execute(
                "SELECT id, title, parent_id FROM sections WHERE parent_id = ?",
                (parent_id,),
            )

        children = self.cursor.fetchall()

        for child in children:
            if child[0] in ids_to_show or child[0] in parents_to_show:
                node = self.tree.insert(parent_node, "end", child[0], text=child[1])
                self.tree.see(node)  # Ensure the node is visible
                self.populate_filtered_tree(
                    child[0], node, ids_to_show, parents_to_show
                )

    def on_closing(self):
        """Handle window closing event."""
        try:
            self.save_data()  # Save any pending changes
            self.conn.close()  # Close the database connection
            self.root.destroy()
        except Exception as e:
            print(f"Error during closing: {e}")
            self.root.destroy()

    def setup_database(self):
        """Create tables if they do not exist."""
        self.cursor.execute(
            """
        CREATE TABLE IF NOT EXISTS sections (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            parent_id INTEGER,
            title TEXT,
            type TEXT, -- 'header', 'category', 'subcategory', or 'subheader'
            questions TEXT, -- JSON array of questions
            placement INTEGER -- Placement of items within the same parent
        )
        """
        )
        self.conn.commit()

        # Add 'placement' column if it doesn't already exist
        try:
            self.cursor.execute("ALTER TABLE sections ADD COLUMN placement INTEGER")
        except sqlite3.OperationalError:
            pass  # Ignore if the column already exists

    def calculate_numbering(self, numbering_dict):
        """Assign hierarchical numbering to tree nodes based on the provided numbering dictionary."""
        for node in self.tree.get_children():
            self.apply_numbering_recursive(node, numbering_dict)

    def apply_numbering_recursive(self, node, numbering_dict):
        """Apply numbering to a node and its children recursively."""
        node_id = self.get_item_id(node)
        if node_id in numbering_dict:
            logical_title = self.tree.item(node, "text").split(". ", 1)[
                -1
            ]  # Remove existing numbering
            display_title = f"{numbering_dict[node_id]}. {logical_title}"
            self.tree.item(node, text=display_title)

        for child in self.tree.get_children(node):
            self.apply_numbering_recursive(child, numbering_dict)

    def generate_numbering(self):
        """Generate a numbering dictionary for all items based on the database hierarchy."""
        numbering_dict = {}

        def recursive_numbering(parent_id=None, prefix=""):
            # Retrieve children based on parent_id
            self.cursor.execute(
                """
                SELECT id, placement FROM sections
                WHERE parent_id IS ?
                ORDER BY placement
            """,
                (parent_id,),
            )
            children = self.cursor.fetchall()

            for idx, (child_id, _) in enumerate(children, start=1):
                number = f"{prefix}{idx}"
                numbering_dict[child_id] = number
                recursive_numbering(child_id, f"{number}.")

        recursive_numbering()  # Start from the root
        return numbering_dict

    def load_from_database(self):
        """Load data from the database and populate the Treeview."""
        try:
            # Clear the treeview
            self.tree.delete(*self.tree.get_children())

            expanded_items = self.get_expanded_items()

            self.cursor.execute(
                """
                SELECT id, title, type, parent_id, placement
                FROM sections 
                ORDER BY placement, id
            """
            )
            sections = self.cursor.fetchall()

            # Populate the treeview
            def populate_tree(parent_id, parent_node):
                current_level = [s for s in sections if s[3] == parent_id]
                for section in current_level:
                    node = self.tree.insert(
                        parent_node, "end", section[0], text=section[1]
                    )
                    populate_tree(section[0], node)

            numbering_dict = self.generate_numbering()  # Generate numbering dictionary
            populate_tree(None, "")
            self.calculate_numbering(numbering_dict)  # Pass only numbering_dict
            self.restore_expansion_state(expanded_items)
        except Exception as e:
            print(f"Error in load_from_database: {e}")

    def add_header(self):
        previous_selection = self.tree.selection()
        title = f"Header {len(self.tree.get_children()) + 1}"
        self.cursor.execute(
            "INSERT INTO sections (title, type, parent_id) VALUES (?, 'header', NULL)",
            (title,),
        )
        self.conn.commit()
        self.load_from_database()
        if previous_selection:
            self.select_item(previous_selection[0])

    def add_category(self):
        previous_selection = self.tree.selection()
        if (
            not previous_selection
            or self.get_item_type(previous_selection[0]) != "header"
        ):
            messagebox.showerror("Error", "Please select a header to add a category.")
            return
        title = f"Category {len(self.tree.get_children(previous_selection[0])) + 1}"
        parent_id = self.get_item_id(previous_selection[0])
        self.cursor.execute(
            "INSERT INTO sections (title, type, parent_id) VALUES (?, 'category', ?)",
            (title, parent_id),
        )
        self.conn.commit()
        self.load_from_database()
        self.select_item(previous_selection[0])

    def add_subcategory(self):
        previous_selection = self.tree.selection()
        if (
            not previous_selection
            or self.get_item_type(previous_selection[0]) != "category"
        ):
            messagebox.showerror(
                "Error", "Please select a category to add a subcategory."
            )
            return
        title = f"Subcategory {len(self.tree.get_children(previous_selection[0])) + 1}"
        parent_id = self.get_item_id(previous_selection[0])
        self.cursor.execute(
            "INSERT INTO sections (title, type, parent_id) VALUES (?, 'subcategory', ?)",
            (title, parent_id),
        )
        self.conn.commit()
        self.load_from_database()
        self.select_item(previous_selection[0])

    def add_subheader(self):
        """Add a subheader (H4) to the selected subcategory (H3)."""
        selected = self.tree.selection()
        if not selected:
            messagebox.showerror(
                "Error", "246: Please select a subcategory to add a subheader."
            )
            return

        selected_id = selected[0]  # Get the selected item ID

        # Get the type directly from database
        query = "SELECT type FROM sections WHERE id = ?"
        self.cursor.execute(query, (int(selected_id),))  # Explicitly convert to int
        result = self.cursor.fetchone()

        if not result or result[0] != "subcategory":
            messagebox.showerror(
                "Error",
                f"260: Please select a subcategory (H3) to add a subheader. Current type: '{result[0] if result else 'None'}'",
            )
            return

        # Get the next placement value
        self.cursor.execute(
            "SELECT COALESCE(MAX(placement), 0) + 1 FROM sections WHERE parent_id = ?",
            (int(selected_id),),  # Explicitly convert to int
        )
        next_placement = self.cursor.fetchone()[0]

        # Create the title and insert the new subheader
        title = f"Sub Header {next_placement}"
        insert_query = "INSERT INTO sections (title, type, parent_id, placement) VALUES (?, ?, ?, ?)"
        params = (title, "subheader", int(selected_id), next_placement)

        self.cursor.execute(insert_query, params)
        self.conn.commit()

        # Save the expanded state and current selection
        expanded_items = self.get_expanded_items()

        # Refresh the tree
        self.load_from_database()

        # Restore expansion state
        self.restore_expansion_state(expanded_items)

        # Reselect the parent H3
        self.select_item(selected_id)

    def save_data(self):
        """Save data for the currently edited item."""
        if self.last_selected_item_id is None:
            return  # No item to save

        logical_title = (
            self.title_entry.get().strip().split(". ", 1)[-1]
        )  # Remove numbering
        questions = json.dumps(self.questions_text.get(1.0, tk.END).strip().split("\n"))

        # Update the database with the logical title
        self.cursor.execute(
            "UPDATE sections SET title = ?, questions = ? WHERE id = ?",
            (logical_title, questions, self.last_selected_item_id),
        )
        self.conn.commit()

        # Update the TreeView display with the logical title
        if self.tree.exists(self.last_selected_item_id):
            self.tree.item(self.last_selected_item_id, text=logical_title)

        numbering_dict = self.generate_numbering()  # Generate numbering dictionary
        self.calculate_numbering(numbering_dict)  # Pass only numbering_dict

    def delete_selected(self):
        """Deletes the selected item and all its children, ensuring parent restrictions."""
        selected = self.tree.selection()
        if not selected:
            messagebox.showerror("Error", "Please select an item to delete.")
            return

        item_id = self.get_item_id(selected[0])
        item_type = self.get_item_type(selected[0])

        # Check if the item has children
        self.cursor.execute(
            "SELECT COUNT(*) FROM sections WHERE parent_id = ?", (item_id,)
        )
        has_children = self.cursor.fetchone()[0] > 0

        if has_children:
            messagebox.showerror(
                "Error", f"Cannot delete {item_type} with child items."
            )
            return

        # Confirm deletion
        confirm = messagebox.askyesno(
            "Confirm Deletion",
            f"Are you sure you want to delete the selected {item_type}?",
        )
        if confirm:
            self.cursor.execute("DELETE FROM sections WHERE id = ?", (item_id,))
            self.conn.commit()

            # Remove the item from the Treeview
            self.tree.delete(selected[0])

            # Reset the editor and last selected item
            self.last_selected_item_id = None
            self.title_entry.delete(0, tk.END)
            self.questions_text.delete(1.0, tk.END)

            print(f"362->Deleted: {item_type.capitalize()} deleted successfully.")

    def reset_database(self):
        """Prompt for a new database file and reset the Treeview."""
        try:
            new_db_path = asksaveasfilename(
                defaultextension=".db",
                filetypes=[("SQLite Database", "*.db")],
                title="Create New Database File",
            )
            if not new_db_path:
                return  # User cancelled
            self.conn.close()  # Close the current database connection

            # Create a new database and reinitialize
            self.conn = sqlite3.connect(new_db_path)
            self.cursor = self.conn.cursor()
            self.setup_database()
            self.initialize_placement()

            # Reset the Treeview
            self.tree.delete(*self.tree.get_children())
            messagebox.showinfo("Success", f"New database created: {new_db_path}")

        except Exception as e:
            messagebox.showerror(
                "Error", f"An error occurred while resetting the database: {e}"
            )

    def load_database_from_file(self):
        """Load an existing database file and update the Treeview."""
        try:
            db_path = askopenfilename(
                filetypes=[("SQLite Database", "*.db")], title="Select Database File"
            )
            if not db_path:
                return  # User cancelled
            self.conn.close()  # Close the current database connection

            # Open the selected database
            self.conn = sqlite3.connect(db_path)
            self.cursor = self.conn.cursor()

            # Clear the treeview
            self.tree.delete(*self.tree.get_children())

            # Reload the Treeview from the new database
            self.load_from_database()
            messagebox.showinfo("Success", f"Database loaded: {db_path}")

        except sqlite3.DatabaseError:
            messagebox.showerror(
                "Error", "The selected file is not a valid SQLite database."
            )
        except Exception as e:
            messagebox.showerror("Error", f"An unexpected error occurred: {e}")

    def get_expanded_items(self):
        """Get a list of expanded items in the Treeview."""
        expanded_items = []
        for item in self.tree.get_children():
            expanded_items.extend(self.get_expanded_items_recursively(item))
        return expanded_items

    def get_expanded_items_recursively(self, item):
        """Recursively check for expanded items."""
        expanded_items = []
        if self.tree.item(item, "open"):
            expanded_items.append(item)
            for child in self.tree.get_children(item):
                expanded_items.extend(self.get_expanded_items_recursively(child))
        return expanded_items

    def restore_expansion_state(self, expanded_items):
        """Restore the expanded state of items in the Treeview."""
        for item in expanded_items:
            self.tree.item(item, open=True)

    def load_from_json(self):
        """Load JSON from file and populate the database with hierarchical data."""
        # Open a file dialog for selecting a JSON file
        file_path = askopenfilename(
            filetypes=[("JSON Files", "*.json")], title="Select JSON File"
        )
        if not file_path:  # If the user cancels the file dialog
            return

        try:
            # Confirm action with the user
            confirm = messagebox.askyesno(
                "Preload Warning",
                "Loading this JSON will populate the database and may cause duplicates. Do you want to continue?",
            )
            if not confirm:
                return

            # Load the JSON data
            with open(file_path, "r") as file:
                data = json.load(file)

            # Validate JSON structure
            self.validate_json_structure(data)

            # Populate the database
            def insert_section(title, section_type, placement, parent_id=None):
                self.cursor.execute(
                    "INSERT INTO sections (title, type, parent_id, placement) VALUES (?, ?, ?, ?)",
                    (title, section_type, parent_id, placement),
                )
                return self.cursor.lastrowid

            for h1_idx, h1_item in enumerate(data.get("h1", []), start=1):
                h1_id = insert_section(h1_item["name"], "header", h1_idx)
                for h2_idx, h2_item in enumerate(h1_item.get("h2", []), start=1):
                    h2_id = insert_section(h2_item["name"], "category", h2_idx, h1_id)
                    for h3_idx, h3_item in enumerate(h2_item.get("h3", []), start=1):
                        insert_section(h3_item["name"], "subcategory", h3_idx, h2_id)

            self.conn.commit()
            self.load_from_database()
            messagebox.showinfo(
                "Success", f"JSON data successfully loaded from {file_path}."
            )

        except FileNotFoundError:
            messagebox.showerror("Error", f"File not found: {file_path}")
        except json.JSONDecodeError:
            messagebox.showerror(
                "Error", "Invalid JSON format. Please select a valid JSON file."
            )
        except ValueError as ve:
            messagebox.showerror("Error", f"Invalid JSON structure: {ve}")
        except Exception as e:
            messagebox.showerror("Error", f"An unexpected error occurred: {e}")

    def validate_json_structure(self, data):
        """Validate the hierarchical structure of the JSON data."""
        if not isinstance(data, dict) or "h1" not in data:
            raise ValueError("Root JSON must be a dictionary with an 'h1' key.")

        for h1_item in data.get("h1", []):
            if not isinstance(h1_item, dict) or "name" not in h1_item:
                raise ValueError(
                    f"Each 'h1' item must be a dictionary with a 'name'. Invalid item: {h1_item}"
                )

            # h2 is optional, but if it exists, validate it
            if "h2" in h1_item:
                if not isinstance(h1_item["h2"], list):
                    raise ValueError(
                        f"'h2' must be a list in 'h1' item. Invalid 'h1' item: {h1_item}"
                    )

                for h2_item in h1_item["h2"]:
                    if not isinstance(h2_item, dict) or "name" not in h2_item:
                        raise ValueError(
                            f"Each 'h2' item must be a dictionary with a 'name'. Invalid item: {h2_item}"
                        )

                    # h3 is optional, but if it exists, validate it
                    if "h3" in h2_item:
                        if not isinstance(h2_item["h3"], list):
                            raise ValueError(
                                f"'h3' must be a list in 'h2' item. Invalid 'h2' item: {h2_item}"
                            )

                        for h3_item in h2_item["h3"]:
                            if not isinstance(h3_item, dict) or "name" not in h3_item:
                                raise ValueError(
                                    f"Each 'h3' item must be a dictionary with a 'name'. Invalid item: {h3_item}"
                                )

    def initialize_placement(self):
        """Assign default placement for existing rows if placement is NULL."""
        try:
            self.cursor.execute(
                """
            WITH RECURSIVE section_hierarchy(id, parent_id, level) AS (
                SELECT id, parent_id, 0 FROM sections WHERE parent_id IS NULL
                UNION ALL
                SELECT s.id, s.parent_id, h.level + 1
                FROM sections s
                INNER JOIN section_hierarchy h ON s.parent_id = h.id
            )
            SELECT id, ROW_NUMBER() OVER (PARTITION BY parent_id ORDER BY id) AS new_placement
            FROM section_hierarchy
            """
            )
            for row in self.cursor.fetchall():
                self.cursor.execute(
                    "SELECT placement FROM sections WHERE id = ?", (row[0],)
                )
                existing_placement = self.cursor.fetchone()[0]
                if existing_placement is None:
                    self.cursor.execute(
                        "UPDATE sections SET placement = ? WHERE id = ?",
                        (row[1], row[0]),
                    )
            self.conn.commit()
        except Exception as e:
            print(f"Error in initialize_placement: {e}")
            self.conn.rollback()

    def swap_placement(self, item_id1, item_id2):
        """Swap the placement of two items in the database and ensure consistency."""
        try:
            # Get current placements
            self.cursor.execute(
                "SELECT placement FROM sections WHERE id = ?", (item_id1,)
            )
            placement1 = self.cursor.fetchone()[0] or 0  # Handle NULL

            self.cursor.execute(
                "SELECT placement FROM sections WHERE id = ?", (item_id2,)
            )
            placement2 = self.cursor.fetchone()[0] or 0  # Handle NULL

            # Perform the swap
            self.cursor.execute(
                "UPDATE sections SET placement = ? WHERE id = ?", (placement2, item_id1)
            )
            self.cursor.execute(
                "UPDATE sections SET placement = ? WHERE id = ?", (placement1, item_id2)
            )

            # Commit changes
            self.conn.commit()

            # Post-commit verification
            self.cursor.execute(
                "SELECT id, placement FROM sections WHERE id IN (?, ?) ORDER BY id",
                (item_id1, item_id2),
            )
            verification = self.cursor.fetchall()
            # print("Post-commit verification of placements:", verification)

            if not verification or len(verification) != 2:
                raise RuntimeError(
                    "Post-commit verification failed: Expected updated rows not found."
                )

            for row in verification:
                if (row[0] == item_id1 and row[1] != placement2) or (
                    row[0] == item_id2 and row[1] != placement1
                ):
                    raise RuntimeError(
                        "Post-commit verification failed: Placements do not match expected values."
                    )

        except sqlite3.OperationalError as e:
            print(f"Database is locked: {e}")
            self.conn.rollback()
        except Exception as e:
            print(f"Error in swap_placement: {e}")
            self.conn.rollback()

    def fix_placement(self, parent_id):
        """Ensure all children of a parent have sequential placement values."""
        try:
            self.cursor.execute(
                "SELECT id FROM sections WHERE parent_id = ? ORDER BY placement NULLS LAST, id",
                (parent_id,),
            )
            children = self.cursor.fetchall()

            for index, (child_id,) in enumerate(children, start=1):
                self.cursor.execute(
                    "UPDATE sections SET placement = ? WHERE id = ?", (index, child_id)
                )

            self.conn.commit()
        except Exception as e:
            print(f"531: Error in fix_placement: {e}")
            self.conn.rollback()

    def move_up(self):
        selected = self.tree.selection()
        if not selected:
            return

        item_id = self.get_item_id(selected[0])
        parent_id = self.tree.parent(selected[0])
        siblings = self.tree.get_children(parent_id)

        index = siblings.index(selected[0])
        if index > 0:  # If not the first item
            self.fix_placement(parent_id)  # Ensure placements are valid
            self.swap_placement(item_id, self.get_item_id(siblings[index - 1]))

            # Refresh the TreeView
            self.refresh_tree()

            self.select_item(selected[0])

    def move_down(self):
        selected = self.tree.selection()
        if not selected:
            return

        item_id = self.get_item_id(selected[0])
        parent_id = self.tree.parent(selected[0])
        siblings = self.tree.get_children(parent_id)

        # Find the index of the selected item
        index = siblings.index(selected[0])
        if index < len(siblings) - 1:  # If not the last item
            self.fix_placement(parent_id)  # Ensure placements are valid
            self.swap_placement(item_id, self.get_item_id(siblings[index + 1]))

            # Save positions and refresh the tree
            self.refresh_tree()

            # Reselect the moved item
            self.select_item(selected[0])

    def refresh_tree(self):
        """Reload the TreeView to reflect database changes."""
        try:
            expanded_items = self.get_expanded_items()
            self.tree.delete(*self.tree.get_children())
            self.load_from_database()
            self.restore_expansion_state(expanded_items)
        except Exception as e:
            print(f"Error in refresh_tree: {e}")

    def get_item_id(self, node):
        """Get the numeric ID from a tree node ID."""
        try:
            return int(node)
        except (ValueError, TypeError):
            print(f"Warning: Invalid node ID format: {node}")
            return None

    def get_item_type(self, node):
        """Fetch the type of the selected node from the database."""
        try:
            item_id = self.get_item_id(node)
            if item_id is None:
                return None

            self.cursor.execute("SELECT type FROM sections WHERE id = ?", (item_id,))
            result = self.cursor.fetchone()
            return result[0] if result else None
        except Exception as e:
            print(f"Error in get_item_type: {e}")
            return None

    def select_item(self, item_id):
        """Select and focus an item in the treeview."""
        try:
            if self.tree.exists(str(item_id)):
                self.tree.selection_set(str(item_id))
                self.tree.focus(str(item_id))
                self.tree.see(str(item_id))  # Ensure the item is visible
        except Exception as e:
            print(f"Error in select_item: {e}")

    def load_selected(self, event):
        """Load the selected item's data into the editor."""
        self.save_data()  # Save previous item's data

        selected = self.tree.selection()
        if not selected:
            return

        item_id = self.get_item_id(selected[0])
        self.last_selected_item_id = item_id  # Track the current selection

        # Retrieve logical title and questions
        self.cursor.execute(
            "SELECT title, questions FROM sections WHERE id = ?", (item_id,)
        )
        row = self.cursor.fetchone()

        if row:
            logical_title, questions = row

            # Clear current values in the editor
            self.title_entry.delete(0, tk.END)
            self.questions_text.delete(1.0, tk.END)

            # Populate editor fields with logical data
            self.title_entry.insert(0, logical_title)
            if questions:
                self.questions_text.insert(tk.END, "\n".join(json.loads(questions)))

    def export_to_docx(self):
        """creates the docx file based on specs defined"""
        try:
            doc = Document()
            self.cursor.execute(
                "SELECT id, title, type, parent_id, questions, placement FROM sections ORDER BY parent_id, placement"
            )
            sections = self.cursor.fetchall()

            # Add Table of Contents Placeholder
            toc_paragraph = doc.add_paragraph("Table of Contents", style="Heading 1")
            toc_paragraph.add_run("\n(TOC will need to be updated in Word)").italic = (
                True
            )
            doc.add_page_break()  # Add page break after TOC

            def add_custom_heading(doc, text, level):
                """Add a custom heading with specific formatting and indentation."""
                paragraph = doc.add_heading(level=level)
                if len(paragraph.runs) == 0:
                    run = paragraph.add_run()
                else:
                    run = paragraph.runs[0]
                run.text = text
                run.font.name = DOC_FONT
                run.bold = True

                # Apply colors and underline based on level
                if level == 1:  # H1 - Brick Red
                    run.font.size = Pt(H1_SIZE)
                    run.font.color.rgb = RGBColor(178, 34, 34)  # Brick red
                elif level == 2:  # H2 - Navy Blue
                    run.font.size = Pt(H2_SIZE)
                    run.font.color.rgb = RGBColor(0, 0, 128)  # Navy blue
                elif level == 3:  # H3 - Black
                    run.font.size = Pt(H3_SIZE)
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Black
                elif level == 4:  # H4 - Underlined
                    run.font.size = Pt(H4_SIZE)
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Black underline
                    run.underline = True

                # Adjust paragraph indentation
                paragraph.paragraph_format.left_indent = Inches(
                    INDENT_SIZE * (level - 1)
                )  # Incremental indentation
                return (
                    paragraph.paragraph_format.left_indent.inches
                )  # Return the calculated indentation

            def add_custom_paragraph(doc, text, style="Normal", indent=0):
                """Add a custom paragraph with specific formatting."""
                paragraph = doc.add_paragraph(text, style=style)
                paragraph.paragraph_format.left_indent = Inches(indent)
                paragraph.paragraph_format.space_after = Pt(P_SIZE)
                if len(paragraph.runs) == 0:
                    run = paragraph.add_run()
                else:
                    run = paragraph.runs[0]
                run.font.name = DOC_FONT
                run.font.size = Pt(P_SIZE)
                return paragraph

            def add_to_doc(parent_id, level, numbering_prefix="", is_first_h1=True):
                """Recursively add sections and their children to the document with hierarchical numbering."""
                children = [s for s in sections if s[3] == parent_id]

                for idx, section in enumerate(children, start=1):
                    # Generate numbering dynamically
                    number = f"{numbering_prefix}{idx}"
                    title_with_number = f"{number}. {section[1]}"

                    # Add page break before H1 (except the first one)
                    if level == 1 and not is_first_h1:
                        doc.add_page_break()
                    if level == 1:
                        is_first_h1 = (
                            False  # Update the flag after processing the first H1
                        )

                    # Add heading with numbering
                    parent_indent = add_custom_heading(doc, title_with_number, level)

                    # Validate and load questions
                    try:
                        questions = json.loads(section[4]) if section[4] else []
                    except json.JSONDecodeError:
                        questions = []

                    # Add content: bullet points for H3/H4, plain paragraphs otherwise
                    if not questions:
                        add_custom_paragraph(
                            doc,
                            "(No questions added yet)",
                            style="Normal",
                            indent=parent_indent + INDENT_SIZE,
                        )
                    else:
                        for question in questions:
                            if level >= 3:  # Add bullet points for H3 and H4
                                add_custom_paragraph(
                                    doc,
                                    question,
                                    style="Normal",
                                    indent=parent_indent + INDENT_SIZE,
                                )
                            else:  # Add plain paragraphs for H1 and H2
                                add_custom_paragraph(
                                    doc,
                                    question,
                                    style="Normal",
                                    indent=parent_indent + INDENT_SIZE,
                                )

                    # Recurse for children
                    add_to_doc(
                        section[0],
                        level + 1,
                        numbering_prefix=f"{number}.",
                        is_first_h1=is_first_h1,
                    )

            # Start adding sections from the root
            add_to_doc(None, 1)

            # Ask the user for a save location
            file_path = asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx")],
                title="Save Document As",
            )
            if not file_path:
                return  # User cancelled the save dialog

            # Save the document
            doc.save(file_path)
            messagebox.showinfo(
                "Exported", f"Document exported successfully to {file_path}."
            )

        except Exception as e:
            messagebox.showerror(
                "Export Failed", f"An error occurred during export:\n{e}"
            )


if __name__ == "__main__":
    root = tk.Tk()
    app = OutLineEditorApp(root)
    root.mainloop()
