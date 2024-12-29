from docx import Document
from docx.shared import Pt, Inches, RGBColor
import json
from config import DOC_FONT, H1_SIZE, H2_SIZE, H3_SIZE, H4_SIZE, P_SIZE, INDENT_SIZE
from tkinter.filedialog import asksaveasfilename
from tkinter import messagebox

from database import DatabaseHandler
from manager_encryption import EncryptionManager

def load_sections_for_export(db_handler: DatabaseHandler, root_id=None):
    """Load sections from database with optional root filtering."""
    db_handler.cursor.execute("""
        WITH RECURSIVE descendants AS (
            SELECT id, title, type, parent_id, questions, placement
            FROM sections
            WHERE id = ? OR (? IS NULL AND parent_id IS NULL)
            UNION ALL
            SELECT s.id, s.title, s.type, s.parent_id, s.questions, s.placement
            FROM sections s
            INNER JOIN descendants d ON s.parent_id = d.id
        )
        SELECT id, title, type, parent_id, questions 
        FROM descendants
        ORDER BY parent_id NULLS FIRST, placement, id
    """, (root_id, root_id))
    
    rows = db_handler.cursor.fetchall()
    decrypted_rows = []
    
    for row in rows:
        decrypted_title = db_handler.decrypt_safely(row[1], f"[Section {row[0]}]")
        decrypted_questions = db_handler.decrypt_safely(row[4], "[]")
        decrypted_rows.append((
            row[0],           # id
            decrypted_title,  # title
            row[2],          # type
            row[3],          # parent_id
            decrypted_questions  # questions
        ))
    
    return decrypted_rows

def export_to_docx(db_handler: DatabaseHandler, root_id=None, file_path=None):
    """Creates the docx file based on specs defined."""
    try:
        if root_id is None:
            confirm = messagebox.askyesno(
                "Full Export Warning",
                "No section selected. This will export the entire document which may take some time. Continue?",
                icon='warning'
            )
            if not confirm:
                return

        sections = load_sections_for_export(db_handler, root_id)
        doc = Document()

        # Add Table of Contents Placeholder
        toc_paragraph = doc.add_paragraph("Table of Contents", style="Heading 1")
        toc_paragraph.add_run("\n(TOC will need to be updated in Word)").italic = True
        doc.add_page_break()

        def add_custom_heading(doc, text, level):
            """Add a custom heading with specific formatting and indentation."""
            paragraph = doc.add_heading(level=min(level, 9))
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.text = text
            run.font.name = DOC_FONT
            run.bold = True

            if level == 1:
                run.font.size = Pt(H1_SIZE)
                run.font.color.rgb = RGBColor(178, 34, 34)
            elif level == 2:
                run.font.size = Pt(H2_SIZE)
                run.font.color.rgb = RGBColor(0, 0, 128)
            elif level == 3:
                run.font.size = Pt(H3_SIZE)
                run.font.color.rgb = RGBColor(0, 0, 0)
            elif level == 4:
                run.font.size = Pt(H4_SIZE)
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.underline = True
            else:
                run.font.size = Pt(H4_SIZE)
                run.font.color.rgb = RGBColor(0, 0, 0)

            indent = INDENT_SIZE * (level - 1)
            paragraph.paragraph_format.left_indent = Inches(indent)
            return indent

        def add_custom_paragraph(doc, text, style="Normal", indent=0):
            paragraph = doc.add_paragraph(text, style=style)
            paragraph.paragraph_format.left_indent = Inches(indent)
            paragraph.paragraph_format.space_after = Pt(P_SIZE)
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.font.name = DOC_FONT
            run.font.size = Pt(P_SIZE)
            return paragraph

        def get_section_level(section_type):
            if section_type == "header":
                return 1
            elif section_type == "category":
                return 2
            elif section_type == "subcategory":
                return 3
            else:  # subheader or anything else
                return 4

        def add_to_doc(parent_id, numbering_prefix="", is_first_h1=True):
            children = [s for s in sections if s[3] == parent_id]

            for idx, section in enumerate(children, start=1):
                section_id, title, section_type, _, questions = section
                number = f"{numbering_prefix}{idx}"
                title_with_number = f"{number}. {title}"

                # Determine level based on section type
                current_level = get_section_level(section_type)
                
                if current_level == 1 and not is_first_h1:
                    doc.add_page_break()
                if current_level == 1:
                    is_first_h1 = False

                parent_indent = add_custom_heading(doc, title_with_number, current_level)

                try:
                    questions_list = json.loads(questions) if questions else []
                except json.JSONDecodeError:
                    questions_list = []

                if not questions_list:
                    add_custom_paragraph(
                        doc,
                        "(No questions added yet)",
                        style="Normal",
                        indent=parent_indent + INDENT_SIZE,
                    )
                else:
                    for question in questions_list:
                        add_custom_paragraph(
                            doc,
                            question,
                            style="Normal",
                            indent=parent_indent + INDENT_SIZE,
                        )

                add_to_doc(
                    section_id,
                    numbering_prefix=f"{number}.",
                    is_first_h1=is_first_h1,
                )

        # If root_id is specified, first export the root node itself
        if root_id:
            root_section = [s for s in sections if s[0] == root_id][0]
            section_id, title, section_type, _, questions = root_section
            
            # Add the root section heading
            parent_indent = add_custom_heading(doc, f"1. {title}", get_section_level(section_type))
            
            # Add root section's questions
            try:
                questions_list = json.loads(questions) if questions else []
            except json.JSONDecodeError:
                questions_list = []

            if not questions_list:
                add_custom_paragraph(
                    doc,
                    "(No questions added yet)",
                    style="Normal",
                    indent=parent_indent + INDENT_SIZE,
                )
            else:
                for question in questions_list:
                    add_custom_paragraph(
                        doc,
                        question,
                        style="Normal",
                        indent=parent_indent + INDENT_SIZE,
                    )
        
        # Then process all children
        add_to_doc(root_id)

        if not file_path:
            file_path = asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx")],
                title="Save Document As",
            )
            if not file_path:
                return

        doc.save(file_path)
        messagebox.showinfo(
            "Exported", f"Document exported successfully to {file_path}."
        )

    except Exception as e:
        messagebox.showerror("Export Failed", f"An error occurred during export:\n{e}")